import base64
import io
import json
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

import pymysql
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from chatbi.repository.db import ensure_table_columns


BASE_DIR = Path(__file__).resolve().parent
REPORT_TEMPLATE_DIR = BASE_DIR / "report_templates"
REPORT_OUTPUT_DIR = BASE_DIR / "report_outputs"
DEFAULT_TEMPLATE_ID = "default-management-report"
DEFAULT_TEMPLATE_FILENAME = "default_management_report_template.docx"
CHINA_GENERAL_TEMPLATE_ID = "china-general-business-report"
CHINA_GENERAL_TEMPLATE_FILENAME = "china_general_business_report_template.docx"
TEMPLATE_PLACEHOLDER_GUIDE = [
    "{{report_title}}",
    "{{executive_summary}}",
    "{{management_summary}}",
    "{{professional_analysis}}",
    "{{strategy_recommendations}}",
    "{{management_actions}}",
    "{{risk_watchouts}}",
    "{{dashboard_snapshot}}",
    "{{detail_table}}",
]

REPORT_TEMPLATE_DDL = """
CREATE TABLE IF NOT EXISTS `report_template` (
    `template_id` VARCHAR(80) NOT NULL COMMENT '模板ID',
    `template_name` VARCHAR(255) NOT NULL COMMENT '模板名称',
    `template_kind` VARCHAR(32) NOT NULL DEFAULT 'custom' COMMENT '模板类型',
    `source_format` VARCHAR(16) NOT NULL DEFAULT 'docx' COMMENT '模板源格式',
    `file_name` VARCHAR(255) NOT NULL COMMENT '原始文件名',
    `file_path` VARCHAR(512) NOT NULL COMMENT '文件路径',
    `template_prompt_text` LONGTEXT NULL COMMENT '模板提示词文本',
    `style_profile_json` LONGTEXT NULL COMMENT '样式画像',
    `validation_summary_json` LONGTEXT NULL COMMENT '模板校验摘要',
    `placeholders_json` LONGTEXT NULL COMMENT '占位符列表',
    `is_default` TINYINT(1) NOT NULL DEFAULT 0 COMMENT '是否默认模板',
    `created_at` DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '创建时间',
    `updated_at` DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP COMMENT '更新时间',
    PRIMARY KEY (`template_id`),
    KEY `idx_report_template_updated_at` (`updated_at`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='ChatBI报告模板表';
"""

REPORT_TEMPLATE_MIGRATIONS = {
    'source_format': "ALTER TABLE `report_template` ADD COLUMN `source_format` VARCHAR(16) NOT NULL DEFAULT 'docx' COMMENT '模板源格式' AFTER `template_kind`",
    'template_prompt_text': "ALTER TABLE `report_template` ADD COLUMN `template_prompt_text` LONGTEXT NULL COMMENT '模板提示词文本' AFTER `file_path`",
    'validation_summary_json': "ALTER TABLE `report_template` ADD COLUMN `validation_summary_json` LONGTEXT NULL COMMENT '模板校验摘要' AFTER `style_profile_json`",
}

REPORT_HISTORY_DDL = """
CREATE TABLE IF NOT EXISTS `report_history` (
    `report_id` VARCHAR(80) NOT NULL COMMENT '报告ID',
    `conversation_id` VARCHAR(80) NOT NULL COMMENT '会话ID',
    `template_id` VARCHAR(80) NOT NULL COMMENT '模板ID',
    `template_name` VARCHAR(255) NOT NULL COMMENT '模板名称',
    `template_kind` VARCHAR(32) NOT NULL COMMENT '模板类型',
    `llm_provider` VARCHAR(32) NOT NULL COMMENT '模型引擎',
    `model_name` VARCHAR(128) NOT NULL COMMENT '模型名称',
    `report_title` VARCHAR(255) NOT NULL COMMENT '报告标题',
    `question` LONGTEXT NULL COMMENT '原始问题',
    `metric_definition` VARCHAR(255) NULL COMMENT '指标定义',
    `metric_description` LONGTEXT NULL COMMENT '指标描述',
    `dimensions_json` LONGTEXT NULL COMMENT '维度列表',
    `metrics_json` LONGTEXT NULL COMMENT '指标列表',
    `row_count` INT NOT NULL DEFAULT 0 COMMENT '结果行数',
    `report_payload_json` LONGTEXT NULL COMMENT '报告内容JSON',
    `latest_result_json` LONGTEXT NULL COMMENT '查询结果快照JSON',
    `file_name` VARCHAR(255) NOT NULL COMMENT '文件名',
    `file_path` VARCHAR(512) NOT NULL COMMENT '文件路径',
    `file_size` BIGINT NOT NULL DEFAULT 0 COMMENT '文件大小',
    `created_at` DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '创建时间',
    PRIMARY KEY (`report_id`),
    KEY `idx_report_history_created_at` (`created_at`),
    KEY `idx_report_history_conversation` (`conversation_id`, `created_at`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='ChatBI报告生成历史表';
"""


def ensure_reporting_runtime(conn: pymysql.connections.Connection) -> None:
    REPORT_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    with conn.cursor() as cursor:
        cursor.execute(REPORT_TEMPLATE_DDL)
        cursor.execute(REPORT_HISTORY_DDL)
        ensure_table_columns(cursor, 'report_template', REPORT_TEMPLATE_MIGRATIONS)
    seed_builtin_templates(conn)

def seed_builtin_templates(conn: pymysql.connections.Connection) -> None:
    builtin_templates = [
        {
            "template_id": DEFAULT_TEMPLATE_ID,
            "template_name": "默认管理层商业分析报告模板",
            "template_kind": "default",
            "file_name": DEFAULT_TEMPLATE_FILENAME,
            "is_default": True,
            "builder": create_default_template_file,
        },
        {
            "template_id": CHINA_GENERAL_TEMPLATE_ID,
            "template_name": "中国通用商业分析报告模板",
            "template_kind": "preset",
            "file_name": CHINA_GENERAL_TEMPLATE_FILENAME,
            "is_default": False,
            "builder": create_china_general_template_file,
        },
    ]
    for template in builtin_templates:
        path = REPORT_TEMPLATE_DIR / template["file_name"]
        if not path.exists():
            template["builder"](path)
        metadata = parse_template_file(path)
        upsert_template_record(
            conn,
            template_id=template["template_id"],
            template_name=template["template_name"],
            template_kind=template["template_kind"],
            file_name=path.name,
            file_path=str(path),
            source_format='docx',
            template_prompt_text=metadata["template_prompt_text"],
            style_profile=metadata["style_profile"],
            validation_summary=metadata["validation_summary"],
            placeholders=metadata["placeholders"],
            is_default=template["is_default"],
        )


def upsert_template_record(
    conn: pymysql.connections.Connection,
    *,
    template_id: str,
    template_name: str,
    template_kind: str,
    file_name: str,
    file_path: str,
    source_format: str,
    template_prompt_text: str,
    style_profile: dict[str, Any],
    validation_summary: dict[str, Any],
    placeholders: list[str],
    is_default: bool,
) -> None:
    with conn.cursor() as cursor:
        cursor.execute(
            """
            INSERT INTO `report_template`
            (`template_id`, `template_name`, `template_kind`, `source_format`, `file_name`, `file_path`,
             `template_prompt_text`, `style_profile_json`, `validation_summary_json`, `placeholders_json`, `is_default`)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                `template_name` = VALUES(`template_name`),
                `template_kind` = VALUES(`template_kind`),
                `source_format` = VALUES(`source_format`),
                `file_name` = VALUES(`file_name`),
                `file_path` = VALUES(`file_path`),
                `template_prompt_text` = VALUES(`template_prompt_text`),
                `style_profile_json` = VALUES(`style_profile_json`),
                `validation_summary_json` = VALUES(`validation_summary_json`),
                `placeholders_json` = VALUES(`placeholders_json`),
                `is_default` = VALUES(`is_default`),
                `updated_at` = NOW()
            """,
            (
                template_id,
                template_name,
                template_kind,
                source_format,
                file_name,
                file_path,
                template_prompt_text,
                json.dumps(style_profile, ensure_ascii=False),
                json.dumps(validation_summary, ensure_ascii=False),
                json.dumps(placeholders, ensure_ascii=False),
                int(is_default),
            ),
        )


def create_default_template_file(path: Path) -> None:
    document = Document()
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("ChatBI 管理层商业分析报告模板")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("用于生成经营概览、看板快照、专业分析与策略建议")

    document.add_paragraph("样式说明：上传自定义模板时，系统会自动解析标题、正文、表格和列表样式。", style="Normal")

    document.add_paragraph("报告占位建议", style="Heading 1")
    document.add_paragraph("以下标记仅作为模板样例，系统会自动识别并提取模板风格：", style="Normal")
    for placeholder in TEMPLATE_PLACEHOLDER_GUIDE:
        document.add_paragraph(placeholder, style="List Bullet")

    document.add_paragraph("章节示例", style="Heading 1")
    document.add_paragraph("一、执行摘要", style="Heading 2")
    document.add_paragraph("建议保留管理层摘要、指标看板、问题诊断、策略建议、行动计划和风险提示等结构。", style="Normal")

    document.add_paragraph("二、看板快照", style="Heading 2")
    document.add_paragraph("可放置业务图表、关键指标表和明细摘要。", style="Normal")

    document.add_paragraph("三、策略建议", style="Heading 2")
    document.add_paragraph("适合沉淀经营策略、经营动作和复盘建议。", style="Normal")

    document.save(path)


def create_china_general_template_file(path: Path) -> None:
    document = Document()
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("中国通用商业分析报告模板")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("适用于经营复盘、月度经营分析、区域/渠道/产品专项汇报")

    document.add_paragraph("模板说明", style="Heading 1")
    document.add_paragraph(
        "该模板适合中国企业常见的经营分析汇报结构，强调管理摘要、经营看板、问题诊断、策略建议和行动计划。",
        style="Normal",
    )

    document.add_paragraph("推荐占位符", style="Heading 1")
    placeholder_blocks = TEMPLATE_PLACEHOLDER_GUIDE + [
        "{{market_environment}}",
        "{{business_overview}}",
        "{{channel_analysis}}",
        "{{regional_analysis}}",
        "{{product_analysis}}",
        "{{problem_diagnosis}}",
        "{{resource_requests}}",
    ]
    for placeholder in placeholder_blocks:
        document.add_paragraph(placeholder, style="List Bullet")

    sections = [
        ("一、管理层摘要", "建议给出一句话结论、经营亮点、主要问题和下一步关注重点。"),
        ("二、经营总览与看板", "建议嵌入关键指标看板、核心趋势图和重点明细摘要。"),
        ("三、市场与行业环境", "补充宏观环境、行业趋势、竞争态势或渠道变化。"),
        ("四、区域/渠道/产品专题分析", "按区域、渠道、产品、用户等维度拆解业绩表现与结构变化。"),
        ("五、问题诊断", "针对增长不及预期、结构失衡、退款波动等问题给出根因分析。"),
        ("六、策略建议", "面向管理层给出可执行、可落地的经营策略。"),
        ("七、重点行动计划", "明确责任部门、节奏安排和预期目标。"),
        ("八、风险与资源需求", "补充风险点、依赖条件和资源诉求。"),
        ("九、附录", "可放置数据口径、样本说明、生成 SQL 和明细摘录。"),
    ]
    for heading, description in sections:
        document.add_paragraph(heading, style="Heading 1")
        document.add_paragraph(description, style="Normal")

    document.save(path)


def parse_template_file(file_path: str | Path) -> dict[str, Any]:
    path = Path(file_path)
    try:
        document = Document(path)
    except (PackageNotFoundError, KeyError, ValueError) as exc:
        raise ValueError('格式检索失败：无法识别为有效的 Word 报告模板，请上传标准 .docx 文件') from exc
    except Exception as exc:  # noqa: BLE001
        raise ValueError('格式检索失败：无法识别为有效的 Word 报告模板，请上传标准 .docx 文件') from exc
    style_names = {style.name for style in document.styles if getattr(style, "name", "")}
    placeholders = extract_placeholders(document)
    template_prompt_text = document_to_markdown_text(document)
    style_profile = {
        "title_style": pick_style_name(style_names, ["Title", "标题"]),
        "subtitle_style": pick_style_name(style_names, ["Subtitle", "副标题", "Normal"]),
        "heading_1_style": pick_style_name(style_names, ["Heading 1", "标题 1"]),
        "heading_2_style": pick_style_name(style_names, ["Heading 2", "标题 2"]),
        "body_style": pick_style_name(style_names, ["Normal", "正文"]),
        "bullet_style": pick_style_name(style_names, ["List Bullet", "项目符号", "Normal"]),
        "table_style": pick_table_style(document),
    }
    validation_summary = validate_docx_template(document, placeholders, template_prompt_text)
    return {
        "style_profile": style_profile,
        "placeholders": placeholders,
        "template_prompt_text": template_prompt_text,
        "validation_summary": validation_summary,
        "file_name": path.name,
    }


def extract_placeholders(document: Document) -> list[str]:
    pattern = re.compile(r"\{\{[^{}]+\}\}")
    placeholders: list[str] = []
    for paragraph in document.paragraphs:
        for match in pattern.findall(paragraph.text or ""):
            if match not in placeholders:
                placeholders.append(match)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for match in pattern.findall(cell.text or ""):
                    if match not in placeholders:
                        placeholders.append(match)
    return placeholders


def pick_style_name(style_names: set[str], candidates: list[str]) -> str:
    lower_map = {name.lower(): name for name in style_names}
    for candidate in candidates:
        if candidate in style_names:
            return candidate
        matched = lower_map.get(candidate.lower())
        if matched:
            return matched
    return next(iter(style_names), "")


def pick_table_style(document: Document) -> str:
    if document.tables and document.tables[0].style and document.tables[0].style.name:
        return document.tables[0].style.name
    style_names = {style.name for style in document.styles if getattr(style, "name", "")}
    return pick_style_name(style_names, ["Table Grid", "网格型"])


def paragraph_to_markdown(paragraph: Any) -> str:
    text = str(paragraph.text or '').strip()
    if not text:
        return ''
    style_name = getattr(getattr(paragraph, 'style', None), 'name', '') or ''
    lower_style = style_name.lower()
    if 'heading 1' in lower_style or style_name.startswith('标题 1'):
        return f'# {text}'
    if 'heading 2' in lower_style or style_name.startswith('标题 2'):
        return f'## {text}'
    if 'heading 3' in lower_style or style_name.startswith('标题 3'):
        return f'### {text}'
    if 'list bullet' in lower_style or '项目符号' in style_name:
        return f'- {text}'
    return text


def document_to_markdown_text(document: Document) -> str:
    lines: list[str] = []
    for paragraph in document.paragraphs:
        line = paragraph_to_markdown(paragraph)
        if line:
            lines.append(line)
    for table in document.tables:
        for row in table.rows:
            row_text = ' | '.join(str(cell.text or '').strip() for cell in row.cells if str(cell.text or '').strip())
            if row_text:
                lines.append(f'| {row_text} |')
    return '\n'.join(lines).strip()


def validate_report_prompt_text(text: str) -> dict[str, Any]:
    normalized = str(text or '').strip()
    if not normalized:
        raise ValueError('文件内容问题：报告模板内容为空，无法作为正常的报告输出模板')
    keyword_groups = ['摘要', '分析', '建议', '策略', '风险', '行动', '结论', '管理层', '看板', '复盘', '经营']
    hit_keywords = [keyword for keyword in keyword_groups if keyword in normalized]
    structural_hits = sum(
        1 for pattern in [r'(^|\n)\s*[一二三四五六七八九十]\s*[、.]', r'(^|\n)\s*#', r'\{\{[^{}]+\}\}', r'(^|\n)\s*[-*]\s+']
        if re.search(pattern, normalized, re.MULTILINE)
    )
    line_count = len([line for line in normalized.splitlines() if line.strip()])
    if len(normalized) < 80 or len(hit_keywords) < 3 or (structural_hits == 0 and line_count < 5):
        raise ValueError('文件内容问题：未识别到完整的商业报告结构，至少应包含摘要/分析/建议/风险等报告要素')
    return {
        'content_length': len(normalized),
        'line_count': line_count,
        'keyword_hits': hit_keywords,
        'structural_hits': structural_hits,
        'recognized': True,
    }


def validate_docx_template(document: Document, placeholders: list[str], markdown_text: str) -> dict[str, Any]:
    validation_summary = validate_report_prompt_text(markdown_text)
    heading_count = sum(
        1 for paragraph in document.paragraphs
        if ('heading' in (getattr(getattr(paragraph, 'style', None), 'name', '') or '').lower()) or (getattr(getattr(paragraph, 'style', None), 'name', '') or '').startswith('标题')
    )
    paragraph_count = len([paragraph for paragraph in document.paragraphs if str(paragraph.text or '').strip()])
    if heading_count == 0 and len(placeholders) < 2 and paragraph_count < 6:
        raise ValueError('格式检索失败：未识别到标题层级、占位符或章节结构，无法作为有效的 Word 报告模板')
    validation_summary.update(
        {
            'heading_count': heading_count,
            'paragraph_count': paragraph_count,
            'placeholder_count': len(placeholders),
            'recognized': True,
        }
    )
    return validation_summary


def parse_text_template_file(file_path: str | Path) -> dict[str, Any]:
    path = Path(file_path)
    try:
        raw_text = path.read_text(encoding='utf-8')
    except UnicodeDecodeError:
        try:
            raw_text = path.read_text(encoding='utf-8-sig')
        except Exception as exc:  # noqa: BLE001
            raise ValueError('文件内容问题：文本模板编码无法识别，请使用 UTF-8 编码') from exc
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f'文件内容问题：文本模板读取失败，{exc}') from exc
    placeholders = re.findall(r'\{\{[^{}]+\}\}', raw_text)
    validation_summary = validate_report_prompt_text(raw_text)
    return {
        'style_profile': {
            'title_style': 'Title',
            'subtitle_style': 'Subtitle',
            'heading_1_style': 'Heading 1',
            'heading_2_style': 'Heading 2',
            'body_style': 'Normal',
            'bullet_style': 'List Bullet',
            'table_style': 'Table Grid',
        },
        'placeholders': list(dict.fromkeys(placeholders)),
        'template_prompt_text': raw_text.strip(),
        'validation_summary': validation_summary,
        'file_name': path.name,
    }


def list_report_templates(conn: pymysql.connections.Connection) -> list[dict[str, Any]]:
    with conn.cursor() as cursor:
        cursor.execute(
            """
            SELECT `template_id`, `template_name`, `template_kind`, `source_format`, `file_name`, `file_path`,
                   `template_prompt_text`, `style_profile_json`, `validation_summary_json`, `placeholders_json`, `is_default`, `updated_at`
            FROM `report_template`
            ORDER BY `is_default` DESC, `updated_at` DESC
            """
        )
        rows = list(cursor.fetchall())

    templates: list[dict[str, Any]] = []
    for row in rows:
        if not Path(row["file_path"]).exists():
            continue
        templates.append(
            {
                "template_id": row["template_id"],
                "template_name": row["template_name"],
                "template_kind": row["template_kind"],
                "source_format": row.get("source_format") or 'docx',
                "file_name": row["file_name"],
                "template_prompt_text": row.get("template_prompt_text") or '',
                "is_default": bool(row["is_default"]),
                "style_profile": safe_json_dict(row.get("style_profile_json")),
                "validation_summary": safe_json_dict(row.get("validation_summary_json")),
                "placeholders": safe_json_list(row.get("placeholders_json")),
                "updated_at": str(row["updated_at"]),
            }
        )
    return templates


def get_report_template(conn: pymysql.connections.Connection, template_id: str | None) -> dict[str, Any]:
    if template_id:
        with conn.cursor() as cursor:
            cursor.execute(
                """
                SELECT `template_id`, `template_name`, `template_kind`, `source_format`, `file_name`, `file_path`,
                       `template_prompt_text`, `style_profile_json`, `validation_summary_json`, `placeholders_json`, `is_default`
                FROM `report_template`
                WHERE `template_id` = %s
                """,
                (template_id,),
            )
            row = cursor.fetchone()
        if row and Path(row["file_path"]).exists():
            return {
                "template_id": row["template_id"],
                "template_name": row["template_name"],
                "template_kind": row["template_kind"],
                "source_format": row.get("source_format") or 'docx',
                "file_name": row["file_name"],
                "file_path": row["file_path"],
                "template_prompt_text": row.get("template_prompt_text") or '',
                "style_profile": safe_json_dict(row.get("style_profile_json")),
                "validation_summary": safe_json_dict(row.get("validation_summary_json")),
                "placeholders": safe_json_list(row.get("placeholders_json")),
                "is_default": bool(row["is_default"]),
            }

    with conn.cursor() as cursor:
        cursor.execute(
            """
            SELECT `template_id`, `template_name`, `template_kind`, `source_format`, `file_name`, `file_path`,
                   `template_prompt_text`, `style_profile_json`, `validation_summary_json`, `placeholders_json`, `is_default`
            FROM `report_template`
            WHERE `is_default` = 1
            ORDER BY `updated_at` DESC
            LIMIT 1
            """
        )
        row = cursor.fetchone()
    if not row:
        raise ValueError("未找到可用报告模板")
    return {
        "template_id": row["template_id"],
        "template_name": row["template_name"],
        "template_kind": row["template_kind"],
        "source_format": row.get("source_format") or 'docx',
        "file_name": row["file_name"],
        "file_path": row["file_path"],
        "template_prompt_text": row.get("template_prompt_text") or '',
        "style_profile": safe_json_dict(row.get("style_profile_json")),
        "validation_summary": safe_json_dict(row.get("validation_summary_json")),
        "placeholders": safe_json_list(row.get("placeholders_json")),
        "is_default": bool(row["is_default"]),
    }


def save_uploaded_template(conn: pymysql.connections.Connection, file_storage: Any) -> dict[str, Any]:
    original_name = str(getattr(file_storage, "filename", "") or "").strip()
    lower_name = original_name.lower()
    if not (lower_name.endswith(".docx") or lower_name.endswith(".txt") or lower_name.endswith(".md")):
        raise ValueError("只支持上传 .docx / .txt / .md 报告模板")

    safe_name = sanitize_filename(original_name)
    template_id = f"tpl_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}"
    target_path = REPORT_TEMPLATE_DIR / f"{template_id}_{safe_name}"
    file_storage.save(target_path)
    try:
        if lower_name.endswith('.docx'):
            metadata = parse_template_file(target_path)
            source_format = 'docx'
        else:
            metadata = parse_text_template_file(target_path)
            source_format = 'txt'
        template_name = target_path.stem.replace("_", " ")
        with conn.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO `report_template`
                (`template_id`, `template_name`, `template_kind`, `source_format`, `file_name`, `file_path`,
                 `template_prompt_text`, `style_profile_json`, `validation_summary_json`, `placeholders_json`, `is_default`)
                VALUES (%s, %s, 'custom', %s, %s, %s, %s, %s, %s, %s, 0)
                """,
                (
                    template_id,
                    template_name,
                    source_format,
                    original_name,
                    str(target_path),
                    metadata["template_prompt_text"],
                    json.dumps(metadata["style_profile"], ensure_ascii=False),
                    json.dumps(metadata["validation_summary"], ensure_ascii=False),
                    json.dumps(metadata["placeholders"], ensure_ascii=False),
                ),
            )
    except Exception:
        if target_path.exists():
            try:
                target_path.unlink()
            except Exception:  # noqa: BLE001
                pass
        raise
    return {
        "template_id": template_id,
        "template_name": template_name,
        "template_kind": "custom",
        "source_format": source_format,
        "file_name": original_name,
        "template_prompt_text": metadata["template_prompt_text"],
        "style_profile": metadata["style_profile"],
        "validation_summary": metadata["validation_summary"],
        "placeholders": metadata["placeholders"],
        "is_default": False,
    }


def set_default_report_template(conn: pymysql.connections.Connection, template_id: str) -> dict[str, Any]:
    template_row = get_report_template(conn, template_id)
    with conn.cursor() as cursor:
        cursor.execute("UPDATE `report_template` SET `is_default` = 0")
        cursor.execute(
            "UPDATE `report_template` SET `is_default` = 1, `updated_at` = NOW() WHERE `template_id` = %s",
            (template_row["template_id"],),
        )
    return get_report_template(conn, template_row["template_id"])


def delete_report_template(conn: pymysql.connections.Connection, template_id: str) -> None:
    with conn.cursor() as cursor:
        cursor.execute(
            """
            SELECT `template_id`, `template_kind`, `file_path`, `is_default`
            FROM `report_template`
            WHERE `template_id` = %s
            """,
            (template_id,),
        )
        row = cursor.fetchone()
        if not row:
            raise ValueError("未找到要删除的报告模板")
        if row["template_kind"] in {"default", "preset"}:
            raise ValueError("预置模板不允许删除")
        if row["is_default"]:
            raise ValueError("默认模板不允许直接删除，请先切换默认模板")
        cursor.execute("DELETE FROM `report_template` WHERE `template_id` = %s", (template_id,))

    template_path = Path(row["file_path"])
    if template_path.exists():
        try:
            template_path.unlink()
        except Exception as exc:  # noqa: BLE001
            raise ValueError(f"模板记录已删除，但清理模板文件失败：{exc}") from exc


def sanitize_filename(name: str) -> str:
    base_name = re.sub(r"[^\w\-.]+", "_", name.strip())
    return base_name or "report_template.docx"


def build_template_markdown_text(template_row: dict[str, Any]) -> str:
    source_format = str(template_row.get('source_format') or 'docx').lower()
    template_prompt_text = str(template_row.get('template_prompt_text') or '').strip()
    if source_format in {'txt', 'md'} and template_prompt_text:
        return template_prompt_text
    file_path = Path(template_row.get('file_path') or '')
    if file_path.suffix.lower() == '.docx' and file_path.exists():
        try:
            return document_to_markdown_text(Document(file_path))
        except Exception:  # noqa: BLE001
            pass
    return template_prompt_text


def export_template_sample_bytes(template_row: dict[str, Any], sample_format: str) -> tuple[bytes, str, str]:
    normalized_format = str(sample_format or 'docx').strip().lower()
    if normalized_format == 'docx':
        file_path = Path(template_row.get('file_path') or '')
        if file_path.suffix.lower() != '.docx':
            document = Document()
            create_text_prompt_template_docx(document, template_row)
            return (
                save_document_to_bytes(document),
                f"{sanitize_filename(Path(template_row.get('file_name') or 'report_template').stem)}.docx",
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        return (
            file_path.read_bytes(),
            template_row.get('file_name') or file_path.name,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

    markdown_text = build_template_markdown_text(template_row) or '# 报告模板\n\n暂无模板内容。'
    file_base = sanitize_filename(Path(template_row.get('file_name') or 'report_template').stem)
    return markdown_text.encode('utf-8'), f'{file_base}.txt', 'text/plain'


def create_text_prompt_template_docx(document: Document, template_row: dict[str, Any]) -> None:
    title = document.add_paragraph(style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(template_row.get('template_name') or '报告模板')
    subtitle = document.add_paragraph(style='Subtitle')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('该模板来自文本提示词，系统已自动生成可编辑的 Word 版本')
    for chunk in split_paragraphs(build_template_markdown_text(template_row)):
        document.add_paragraph(chunk, style='Normal')


def build_report_output_name(report_id: str, download_name: str) -> str:
    safe_name = sanitize_filename(download_name)
    return f"{report_id}_{safe_name}"


def safe_json_dict(raw_value: Any) -> dict[str, Any]:
    if isinstance(raw_value, dict):
        return raw_value
    if not raw_value:
        return {}
    try:
        payload = json.loads(str(raw_value))
    except Exception:  # noqa: BLE001
        return {}
    return payload if isinstance(payload, dict) else {}


def safe_json_list(raw_value: Any) -> list[Any]:
    if isinstance(raw_value, list):
        return raw_value
    if not raw_value:
        return []
    try:
        payload = json.loads(str(raw_value))
    except Exception:  # noqa: BLE001
        return []
    return payload if isinstance(payload, list) else []


def build_csv_bytes(latest_result: dict[str, Any]) -> bytes:
    import csv

    buffer = io.StringIO()
    rows = latest_result.get("rows", []) or []
    columns = latest_result.get("columns", []) or []
    if not columns and rows:
        columns = list(rows[0].keys())

    csv_writer = csv.writer(buffer)
    if columns:
        csv_writer.writerow(columns)
        for row in rows:
            csv_writer.writerow([row.get(column, "") for column in columns])
    return buffer.getvalue().encode("utf-8-sig")


def build_chart_word_bytes(
    template_row: dict[str, Any],
    latest_result: dict[str, Any],
    chart_images: list[dict[str, Any]],
) -> bytes:
    document = prepare_document(template_row)
    styles = template_row.get("style_profile", {})

    add_title(document, styles, latest_result.get("chart_title") or latest_result.get("metric_definition") or "图表快照导出")
    add_subtitle(
        document,
        styles,
        f"指标定义：{latest_result.get('metric_definition', '--')} | 指标：{', '.join(latest_result.get('metrics', [])) or '--'}",
    )

    add_heading(document, styles, "看板概览", level=1)
    dashboard_pairs = [
        ("问题", latest_result.get("question", "--")),
        ("指标定义", latest_result.get("metric_definition", "--")),
        ("指标描述", latest_result.get("metric_description", "--")),
        ("维度", "、".join(latest_result.get("dimensions", [])) or "整体汇总"),
        ("指标", "、".join(latest_result.get("metrics", [])) or "--"),
        ("返回行数", str(latest_result.get("row_count", 0))),
    ]
    add_key_value_table(document, dashboard_pairs, styles)

    add_heading(document, styles, "图表快照", level=1)
    if chart_images:
        add_chart_snapshots(document, chart_images, styles)
    else:
        add_body_paragraph(document, styles, "当前结果不适合生成柱图或饼图，因此本次导出仅包含明细结果。")

    add_heading(document, styles, "结果明细", level=1)
    add_result_table(document, latest_result.get("columns", []), latest_result.get("rows", []), styles, max_rows=80)

    add_heading(document, styles, "生成 SQL", level=1)
    add_body_paragraph(document, styles, latest_result.get("generated_sql", "--"))
    return save_document_to_bytes(document)


def build_management_report_docx(
    template_row: dict[str, Any],
    latest_result: dict[str, Any],
    report_payload: dict[str, Any],
    chart_images: list[dict[str, Any]],
) -> bytes:
    document = prepare_document(template_row)
    styles = template_row.get("style_profile", {})

    add_title(document, styles, report_payload.get("report_title") or latest_result.get("metric_definition") or "商业分析报告")
    add_subtitle(
        document,
        styles,
        report_payload.get("report_subtitle")
        or f"问题：{latest_result.get('question', '--')} | 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    )

    add_heading(document, styles, "一、执行摘要", level=1)
    add_body_paragraph(document, styles, report_payload.get("executive_summary", ""))

    add_heading(document, styles, "二、管理层结论", level=1)
    add_body_paragraph(document, styles, report_payload.get("management_summary", ""))

    add_heading(document, styles, "三、看板快照", level=1)
    dashboard_pairs = [
        ("指标定义", latest_result.get("metric_definition", "--")),
        ("指标描述", latest_result.get("metric_description", "--")),
        ("维度", "、".join(latest_result.get("dimensions", [])) or "整体汇总"),
        ("指标", "、".join(latest_result.get("metrics", [])) or "--"),
        ("数据行数", str(latest_result.get("row_count", 0))),
        ("图表标题", latest_result.get("chart_title", "--")),
    ]
    add_key_value_table(document, dashboard_pairs, styles)
    add_chart_snapshots(document, chart_images, styles)

    add_heading(document, styles, "四、关键发现", level=1)
    add_bullet_list(document, report_payload.get("key_findings", []), styles)

    add_heading(document, styles, "五、专业分析", level=1)
    for section in report_payload.get("professional_analysis", []):
        add_heading(document, styles, section.get("title", "分析章节"), level=2)
        add_body_paragraph(document, styles, section.get("content", ""))

    add_heading(document, styles, "六、策略建议", level=1)
    add_bullet_list(document, report_payload.get("strategy_recommendations", []), styles)

    add_heading(document, styles, "七、行动计划", level=1)
    add_bullet_list(document, report_payload.get("management_actions", []), styles)

    add_heading(document, styles, "八、风险与关注点", level=1)
    add_bullet_list(document, report_payload.get("risk_watchouts", []), styles)

    add_heading(document, styles, "九、结果明细摘录", level=1)
    add_result_table(document, latest_result.get("columns", []), latest_result.get("rows", []), styles, max_rows=30)

    add_heading(document, styles, "十、附录", level=1)
    add_body_paragraph(document, styles, report_payload.get("appendix_note", "本报告由 ChatBI 自动生成，建议结合业务背景进行最终审阅。"))
    add_body_paragraph(document, styles, f"生成 SQL：{latest_result.get('generated_sql', '--')}")
    return save_document_to_bytes(document)


def prepare_document(template_row: dict[str, Any]) -> Document:
    path = Path(template_row["file_path"])
    if path.suffix.lower() != '.docx' or not path.exists():
        document = Document()
        return document
    document = Document(path)
    clear_document_body(document)
    return document


def clear_document_body(document: Document) -> None:
    body = document._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if sect_pr is not None and child == sect_pr:
            continue
        body.remove(child)


def save_document_to_bytes(document: Document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_title(document: Document, styles: dict[str, Any], text: str) -> None:
    paragraph = document.add_paragraph(style=resolve_style(document, styles.get("title_style")))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text or "商业分析报告")
    run.bold = True
    run.font.size = Pt(20)


def add_subtitle(document: Document, styles: dict[str, Any], text: str) -> None:
    paragraph = document.add_paragraph(style=resolve_style(document, styles.get("subtitle_style")))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text or "")


def add_heading(document: Document, styles: dict[str, Any], text: str, level: int = 1) -> None:
    style_name = styles.get("heading_1_style") if level == 1 else styles.get("heading_2_style")
    document.add_paragraph(text or "", style=resolve_style(document, style_name))


def add_body_paragraph(document: Document, styles: dict[str, Any], text: str) -> None:
    for chunk in split_paragraphs(text):
        paragraph = document.add_paragraph(style=resolve_style(document, styles.get("body_style")))
        paragraph.add_run(chunk)


def add_bullet_list(document: Document, items: list[str], styles: dict[str, Any]) -> None:
    if not items:
        add_body_paragraph(document, styles, "暂无补充建议。")
        return
    for item in items:
        document.add_paragraph(item, style=resolve_style(document, styles.get("bullet_style")))


def add_key_value_table(document: Document, pairs: list[tuple[str, str]], styles: dict[str, Any]) -> None:
    table = document.add_table(rows=1, cols=2)
    apply_table_style(table, styles)
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "内容"
    for key, value in pairs:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)


def add_result_table(
    document: Document,
    columns: list[str],
    rows: list[dict[str, Any]],
    styles: dict[str, Any],
    max_rows: int = 30,
) -> None:
    if not columns:
        add_body_paragraph(document, styles, "当前结果无可展示数据。")
        return
    table = document.add_table(rows=1, cols=len(columns))
    apply_table_style(table, styles)
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = str(column)
    for row_data in rows[:max_rows]:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = str(row_data.get(column, ""))
    if len(rows) > max_rows:
        add_body_paragraph(document, styles, f"本报告仅展示前 {max_rows} 行数据，完整明细建议通过 CSV 导出。")


def add_chart_snapshots(document: Document, chart_images: list[dict[str, Any]], styles: dict[str, Any]) -> None:
    if not chart_images:
        add_body_paragraph(document, styles, "当前结果未生成可嵌入的图表快照。")
        return
    for chart in chart_images:
        add_heading(document, styles, chart.get("title", "图表快照"), level=2)
        image_bytes = decode_data_url(chart.get("png_data_url", ""))
        if not image_bytes:
            add_body_paragraph(document, styles, "图表图片生成失败，已跳过。")
            continue
        picture_stream = io.BytesIO(image_bytes)
        document.add_picture(picture_stream, width=Inches(6.5))
        if chart.get("caption"):
            add_body_paragraph(document, styles, chart["caption"])


def apply_table_style(table: Any, styles: dict[str, Any]) -> None:
    style_name = styles.get("table_style")
    if style_name:
        try:
            table.style = style_name
        except Exception:  # noqa: BLE001
            try:
                table.style = "Table Grid"
            except Exception:  # noqa: BLE001
                pass


def resolve_style(document: Document, style_name: str | None) -> str | None:
    if not style_name:
        return None
    try:
        document.styles[style_name]
        return style_name
    except Exception:  # noqa: BLE001
        return None


def split_paragraphs(text: str) -> list[str]:
    cleaned = str(text or "").strip()
    if not cleaned:
        return [""]
    return [item.strip() for item in re.split(r"\n{2,}", cleaned) if item.strip()]


def decode_data_url(data_url: str) -> bytes:
    if not data_url or "," not in data_url:
        return b""
    try:
        return base64.b64decode(data_url.split(",", 1)[1])
    except Exception:  # noqa: BLE001
        return b""


def save_report_history(
    conn: pymysql.connections.Connection,
    *,
    conversation_id: str,
    template_row: dict[str, Any],
    latest_result: dict[str, Any],
    report_payload: dict[str, Any],
    llm_provider: str,
    model_name: str,
    document_bytes: bytes,
    download_name: str,
) -> dict[str, Any]:
    report_id = f"rpt_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}"
    output_name = build_report_output_name(report_id, download_name)
    output_path = REPORT_OUTPUT_DIR / output_name
    output_path.write_bytes(document_bytes)

    with conn.cursor() as cursor:
        cursor.execute(
            """
            INSERT INTO `report_history`
            (`report_id`, `conversation_id`, `template_id`, `template_name`, `template_kind`,
             `llm_provider`, `model_name`, `report_title`, `question`, `metric_definition`,
             `metric_description`, `dimensions_json`, `metrics_json`, `row_count`,
             `report_payload_json`, `latest_result_json`, `file_name`, `file_path`, `file_size`)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """,
            (
                report_id,
                str(conversation_id or ""),
                str(template_row.get("template_id") or ""),
                str(template_row.get("template_name") or ""),
                str(template_row.get("template_kind") or ""),
                str(llm_provider or ""),
                str(model_name or ""),
                str(report_payload.get("report_title") or latest_result.get("metric_definition") or "商业分析报告"),
                str(latest_result.get("question") or ""),
                str(latest_result.get("metric_definition") or ""),
                str(latest_result.get("metric_description") or ""),
                json.dumps(latest_result.get("dimensions", []), ensure_ascii=False),
                json.dumps(latest_result.get("metrics", []), ensure_ascii=False),
                int(latest_result.get("row_count") or 0),
                json.dumps(report_payload, ensure_ascii=False, default=str),
                json.dumps(latest_result, ensure_ascii=False, default=str),
                download_name,
                str(output_path),
                len(document_bytes),
            ),
        )

    return get_report_history_detail(conn, report_id)


def list_report_history(conn: pymysql.connections.Connection, limit: int = 200) -> list[dict[str, Any]]:
    with conn.cursor() as cursor:
        cursor.execute(
            """
            SELECT `report_id`, `conversation_id`, `template_id`, `template_name`, `template_kind`,
                   `llm_provider`, `model_name`, `report_title`, `question`, `metric_definition`,
                   `row_count`, `file_name`, `file_path`, `file_size`, `created_at`
            FROM `report_history`
            ORDER BY `created_at` DESC, `report_id` DESC
            LIMIT %s
            """,
            (int(limit),),
        )
        rows = list(cursor.fetchall())
    return [_normalize_report_history_summary(row) for row in rows]


def get_report_history_detail(conn: pymysql.connections.Connection, report_id: str) -> dict[str, Any]:
    with conn.cursor() as cursor:
        cursor.execute(
            """
            SELECT `report_id`, `conversation_id`, `template_id`, `template_name`, `template_kind`,
                   `llm_provider`, `model_name`, `report_title`, `question`, `metric_definition`,
                   `metric_description`, `dimensions_json`, `metrics_json`, `row_count`,
                   `report_payload_json`, `latest_result_json`, `file_name`, `file_path`,
                   `file_size`, `created_at`
            FROM `report_history`
            WHERE `report_id` = %s
            """,
            (report_id,),
        )
        row = cursor.fetchone()
    if not row:
        raise ValueError("未找到对应的报告历史")
    detail = _normalize_report_history_summary(row)
    detail["metric_description"] = str(row.get("metric_description") or "")
    detail["report_payload"] = safe_json_dict(row.get("report_payload_json"))
    detail["latest_result"] = safe_json_dict(row.get("latest_result_json"))
    return detail


def get_report_history_file(conn: pymysql.connections.Connection, report_id: str) -> dict[str, Any]:
    detail = get_report_history_detail(conn, report_id)
    file_path = Path(detail["file_path"])
    if not file_path.exists():
        raise ValueError("报告文件不存在，可能已被清理")
    return detail


def _normalize_report_history_summary(row: dict[str, Any]) -> dict[str, Any]:
    return {
        "report_id": row["report_id"],
        "conversation_id": row["conversation_id"],
        "template_id": row["template_id"],
        "template_name": row["template_name"],
        "template_kind": row["template_kind"],
        "llm_provider": row["llm_provider"],
        "model_name": row["model_name"],
        "report_title": row["report_title"],
        "question": row.get("question") or "",
        "metric_definition": row.get("metric_definition") or "",
        "dimensions": safe_json_list(row.get("dimensions_json")),
        "metrics": safe_json_list(row.get("metrics_json")),
        "row_count": int(row.get("row_count") or 0),
        "file_name": row["file_name"],
        "file_path": row["file_path"],
        "file_size": int(row.get("file_size") or 0),
        "file_exists": Path(row["file_path"]).exists(),
        "created_at": str(row["created_at"]),
    }
